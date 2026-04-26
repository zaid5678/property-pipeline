"""
Document automation — generates NDA and Sourcing Agreement as .docx files
using python-docx.  No template files needed — everything is built in code.

Usage:
  from documents.generator import generate_nda, generate_sourcing_agreement
"""

import logging
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path(__file__).parent.parent / "output"


def _heading(doc, text: str, level: int = 1):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
    return p


def _body(doc, text: str):
    from docx.shared import Pt
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def _clause(doc, number: str, title: str, body: str):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(11)

    p2 = doc.add_paragraph(body)
    p2.paragraph_format.left_indent = Pt(18)
    for run in p2.runs:
        run.font.size = Pt(11)


def generate_nda(
    investor_name: str,
    property_ref: str,
    sourcer_name: str = "Your Name",
    sourcer_company: str = "Your Company",
    date: str | None = None,
) -> Path:
    """
    Generate a Confidentiality Agreement (NDA) as a .docx file.
    Returns the path to the generated file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("Run: pip install python-docx")

    OUTPUT_DIR.mkdir(exist_ok=True)
    date_str = date or datetime.now().strftime("%d %B %Y")
    filename = OUTPUT_DIR / f"NDA_{investor_name.replace(' ','_')}_{property_ref}.docx"

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CONFIDENTIALITY AGREEMENT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

    doc.add_paragraph()

    # Parties
    _body(doc, (
        f"This Confidentiality Agreement (\"Agreement\") is entered into on {date_str} "
        f"between:"
    ))
    doc.add_paragraph()
    _body(doc, f"DISCLOSING PARTY: {sourcer_name}, trading as {sourcer_company} "
               f"(\"the Sourcer\")")
    _body(doc, f"RECEIVING PARTY:  {investor_name} (\"the Investor\")")
    doc.add_paragraph()
    _body(doc, f"PROPERTY REFERENCE: {property_ref}")
    doc.add_paragraph()

    _heading(doc, "1. Purpose", level=2)
    _body(doc, (
        "The Sourcer intends to disclose certain confidential information relating to "
        "an off-market property investment opportunity (the \"Property\") to the Investor "
        "for the sole purpose of the Investor evaluating whether to proceed with the "
        "purchase of the Property."
    ))

    _heading(doc, "2. Confidential Information", level=2)
    _body(doc, (
        "\"Confidential Information\" means all information disclosed by the Sourcer to "
        "the Investor relating to the Property, including but not limited to: the property "
        "address, purchase price, vendor details, survey results, comparable evidence, "
        "rental income estimates, and all financial analysis."
    ))

    _heading(doc, "3. Obligations", level=2)
    _body(doc, "The Investor agrees to:")
    for item in [
        "Keep all Confidential Information strictly confidential and not disclose it to any third party without the Sourcer's prior written consent.",
        "Use the Confidential Information solely to evaluate the Property and not for any other purpose.",
        "Not approach the vendor, estate agent, or any party connected to the Property directly without the Sourcer's involvement.",
        "Not attempt to circumvent the Sourcer's role in the transaction.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(11)

    _heading(doc, "4. Sourcing Fee", level=2)
    _body(doc, (
        "The Investor acknowledges that a sourcing fee is payable to the Sourcer in "
        "accordance with any Sourcing Agreement entered into between the parties. "
        "The obligation to pay the sourcing fee is not affected by the terms of this Agreement."
    ))

    _heading(doc, "5. Duration", level=2)
    _body(doc, (
        "This Agreement shall remain in force for a period of two (2) years from the "
        "date of signing, or until the Property has been purchased or the Investor "
        "has confirmed in writing they are not proceeding, whichever is earlier."
    ))

    _heading(doc, "6. Governing Law", level=2)
    _body(doc, (
        "This Agreement shall be governed by and construed in accordance with the "
        "laws of England and Wales."
    ))

    doc.add_paragraph()
    _body(doc, "IN WITNESS WHEREOF, the parties have executed this Agreement as of the date written above.")
    doc.add_paragraph()

    # Signature blocks
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = "Table Grid"
    headers = sig_table.rows[0].cells
    headers[0].text = "SOURCER"
    headers[1].text = "INVESTOR"
    names = sig_table.rows[1].cells
    names[0].text = f"Name: {sourcer_name}"
    names[1].text = f"Name: {investor_name}"
    sigs = sig_table.rows[2].cells
    sigs[0].text = f"Signature: ___________________\n\nDate: ___________________"
    sigs[1].text = f"Signature: ___________________\n\nDate: ___________________"

    doc.save(str(filename))
    logger.info("[Docs] NDA saved: %s", filename)
    print(f"[Docs] NDA generated: {filename}")
    return filename


def generate_sourcing_agreement(
    investor_name: str,
    investor_email: str,
    property_address: str,
    property_ref: str,
    purchase_price: int,
    sourcing_fee: int = 3000,
    payment_terms: str = "50% on exchange of contracts, 50% on completion",
    sourcer_name: str = "Your Name",
    sourcer_company: str = "Your Company",
    sourcer_email: str = "",
    date: str | None = None,
) -> Path:
    """
    Generate a Sourcing Agreement as a .docx file.
    Returns the path to the generated file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("Run: pip install python-docx")

    OUTPUT_DIR.mkdir(exist_ok=True)
    date_str = date or datetime.now().strftime("%d %B %Y")
    filename = OUTPUT_DIR / f"SourcingAgreement_{investor_name.replace(' ','_')}_{property_ref}.docx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROPERTY SOURCING AGREEMENT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

    doc.add_paragraph()
    _body(doc, f"Date: {date_str}")
    doc.add_paragraph()

    _body(doc, "This Property Sourcing Agreement is made between:")
    doc.add_paragraph()
    _body(doc, f"SOURCER: {sourcer_name}, trading as {sourcer_company} "
               + (f"({sourcer_email})" if sourcer_email else ""))
    _body(doc, f"CLIENT (INVESTOR): {investor_name} ({investor_email})")
    doc.add_paragraph()

    # Property details box
    prop_table = doc.add_table(rows=3, cols=2)
    prop_table.style = "Table Grid"
    prop_table.rows[0].cells[0].text = "Property Address"
    prop_table.rows[0].cells[1].text = property_address
    prop_table.rows[1].cells[0].text = "Property Reference"
    prop_table.rows[1].cells[1].text = property_ref
    prop_table.rows[2].cells[0].text = "Agreed Purchase Price"
    prop_table.rows[2].cells[1].text = f"£{purchase_price:,}"
    doc.add_paragraph()

    _heading(doc, "1. Services", level=2)
    _body(doc, (
        "The Sourcer agrees to introduce the Client to the above Property, which has "
        "been sourced by the Sourcer on an off-market basis. The Sourcer will provide: "
        "property details, comparable evidence, rental income estimates, and an "
        "introduction to the vendor or their representative."
    ))

    _heading(doc, "2. Sourcing Fee", level=2)
    _body(doc, (
        f"In consideration of the services provided, the Client agrees to pay the "
        f"Sourcer a sourcing fee of £{sourcing_fee:,} (the \"Fee\")."
    ))
    _body(doc, f"Payment terms: {payment_terms}.")
    _body(doc, (
        "The Fee is payable regardless of whether the Client uses any third-party "
        "finance. The Fee is non-refundable once exchange of contracts has taken place."
    ))

    _heading(doc, "3. Client Obligations", level=2)
    for item in [
        "Conduct their own due diligence on the Property.",
        "Not approach the vendor directly, bypassing the Sourcer.",
        "Confirm in writing within 5 working days whether they wish to proceed.",
        "Sign a Confidentiality Agreement before receiving full property details.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(11)

    _heading(doc, "4. Compliance", level=2)
    _body(doc, (
        "The Sourcer confirms that this activity is carried out in compliance with "
        "the Estate Agents Act 1979 and the Property Ombudsman Code of Practice. "
        "The Sourcer is not an authorised financial adviser and nothing in this "
        "agreement constitutes financial advice."
    ))

    _heading(doc, "5. Limitation of Liability", level=2)
    _body(doc, (
        "The Sourcer's liability in connection with this Agreement shall not exceed "
        "the amount of the Fee paid. The Sourcer gives no warranty as to the "
        "accuracy of rental income estimates or market value assessments, which are "
        "estimates only."
    ))

    _heading(doc, "6. Governing Law", level=2)
    _body(doc, (
        "This Agreement is governed by the laws of England and Wales. Any disputes "
        "shall be subject to the exclusive jurisdiction of the courts of England and Wales."
    ))

    doc.add_paragraph()
    _body(doc, "Signed and agreed by the parties:")
    doc.add_paragraph()

    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = "Table Grid"
    sig_table.rows[0].cells[0].text = "SOURCER"
    sig_table.rows[0].cells[1].text = "CLIENT / INVESTOR"
    sig_table.rows[1].cells[0].text = f"Name: {sourcer_name}"
    sig_table.rows[1].cells[1].text = f"Name: {investor_name}"
    sig_table.rows[2].cells[0].text = "Signature: ___________________\n\nDate: ___________________"
    sig_table.rows[2].cells[1].text = "Signature: ___________________\n\nDate: ___________________"

    doc.save(str(filename))
    logger.info("[Docs] Sourcing Agreement saved: %s", filename)
    print(f"[Docs] Sourcing Agreement generated: {filename}")
    return filename


def generate_docs_for_investor(
    investor: dict,
    deal: dict,
    sourcer_name: str,
    sourcer_company: str,
    sourcer_email: str,
    sourcing_fee: int,
    payment_terms: str,
) -> dict[str, Path]:
    """Generate both NDA and Sourcing Agreement for an investor/deal pair."""
    from datetime import datetime
    property_ref = f"DP-{datetime.now().strftime('%Y%m%d')}-{deal.get('id','000')}"

    nda_path = generate_nda(
        investor_name=investor["name"],
        property_ref=property_ref,
        sourcer_name=sourcer_name,
        sourcer_company=sourcer_company,
    )

    sa_path = generate_sourcing_agreement(
        investor_name=investor["name"],
        investor_email=investor["email"],
        property_address=deal.get("address", "TBC"),
        property_ref=property_ref,
        purchase_price=deal.get("purchase_price", 0),
        sourcing_fee=sourcing_fee,
        payment_terms=payment_terms,
        sourcer_name=sourcer_name,
        sourcer_company=sourcer_company,
        sourcer_email=sourcer_email,
    )

    return {"nda": nda_path, "sourcing_agreement": sa_path}
